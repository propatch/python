# This is a sample Python script.
from docx import Document
import zipfile

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.


def print_hi(name):
    # Use a breakpoint in the code line below to debug your script.
    print(f'Hi, {name}')  # Press Ctrl+F8 to toggle the breakpoint.


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    print_hi('PyCharm')

document = Document()
z = zipfile.ZipFile('test.docx')
print(z.read('word/document.xml'))
'xml' in z.read('word/document.xml')

document = docx.Document("example.docx")

for paragraph in document.paragraphs:
    paragraph.text = paragraph.text.replace("old", "new")

document.save("example.docx")


z.close()
document.save('test.docx')
